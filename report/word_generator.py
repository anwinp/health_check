

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches, Cm
import os
from django.conf import settings


class WordReportGenerator:
    def __init__(self, output_path):
        self.template_path = os.path.join(settings.BASE_DIR, 'report', 'health_check_report_template.docx')
        print(f"Template path: {self.template_path}")  # Debugging statement
        if not os.path.exists(self.template_path):
            raise FileNotFoundError(f"Template file does not exist at path: {self.template_path}")
        self.document = Document(self.template_path)
        self.output_path = output_path

    def set_font(self, run, font_name, size, bold=False, underline=False):
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.size = Pt(size)
        run.bold = bold
        if underline:
            run.underline = True

    def add_table_from_data(self, node, data):
        if not data:
            print(f"No data available for node {node}")
            return
        

        # Add a node header
        node_paragraph = self.document.add_paragraph(node)
        self.set_font(node_paragraph.runs[0], 'Segoe UI', 14, bold=True)

        # Add table to the document
        table = self.document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'

        # Add the header row
        hdr_cells = table.rows[0].cells
        headers = ["Info Type", "Data"]
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            self.set_font(hdr_cells[i].paragraphs[0].runs[0], 'Segoe UI', 12, bold=True)

        # Add the data rows
        for row in data:
            row_cells = table.add_row().cells
            row_cells[0].text = row["Info Type"]
            row_cells[1].text = row["Data"]
            self.set_font(row_cells[0].paragraphs[0].runs[0], 'Segoe UI', 12)
            self.set_font(row_cells[1].paragraphs[0].runs[0], 'Segoe UI', 12)


    def generate_report(self, data_by_node):
        for node, data in data_by_node.items():
            self.add_table_from_data(node, data)
        self.document.save(self.output_path)

    def get_filepath(self):
        return self.output_path